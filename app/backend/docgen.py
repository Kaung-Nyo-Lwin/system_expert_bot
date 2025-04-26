from docx import Document
import uuid, os
def create_doc(user_input, response):
    filename = f"{uuid.uuid4().hex}.docx"
    doc = Document()
    doc.add_heading("SoftwareDocBot Output", 0)
    doc.add_paragraph(f"User input: {user_input}")
    doc.add_paragraph(f"Bot response: {response}")
    path = os.path.join("static/docs", filename)
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    return filename
